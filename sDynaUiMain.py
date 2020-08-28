
#--------------------Library-------------------#
#----------------------------------------------#
import sys
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import *
from sDynaUi import *
from about import *
from manual import *
import xlsxwriter
import xlrd
from MDOF import *
import matplotlib.pyplot as plt
import numpy as np
import os

#--------------Create Application--------------#
#----------------------------------------------#
Application = QApplication(sys.argv)
WinMain = QMainWindow()
WinMain.setWindowIcon
ui = Ui_sDyna() #sDynaUi.py class isminden kopyalandı.
ui.setupUi(WinMain) #tasarımdaki form ile pencereyi birleştir
WinMain.show() #pencereyi göster.

WinAbout=QDialog()
uiAbout = Ui_Dialog()
uiAbout.setupUi(WinAbout)

WinManual=QDialog()
uiManual = Ui_Manual()
uiManual.setupUi(WinManual)

#---------------Create DataBase----------------#
#----------------------------------------------#
import sqlite3
global curs
global conn

conn = sqlite3.connect("sDynaDB.db")
curs = conn.cursor()

curs.execute("DELETE FROM sDyna")
conn.commit()

curs.execute("CREATE TABLE IF NOT EXISTS sDyna(                     \
                Floor NUMERIC NOT NULL PRIMARY KEY,                             \
                Mass NUMERIC NOT NULL,                              \
                Rigidity NUMERIC NOT NULL)")
conn.commit()

#---------------lineEdit Enable----------------#
#----------------------------------------------#

ui.lne_Mass.setEnabled(False)
ui.lne_Rigidity.setEnabled(False)
ui.progressBar.hide()

def comboact():
    ui.lne_Mass.setEnabled(True)
    ui.lne_Rigidity.setEnabled(True)


#---------------SAVE---------------#
#----------------------------------#
def addData():
    checkUnique = 1
    _lne_Mass = ui.lne_Mass.text()
    _lne_Rigidity = ui.lne_Rigidity.text()
    _cm_Floor = ui.cm_Floor.currentText()
    try:
        _lne_Mass= float(_lne_Mass)
        _lne_Rigidity=float(_lne_Rigidity)
        if bool(_lne_Mass) and bool(_lne_Rigidity) and bool(_cm_Floor)==True:
        
            curs.execute("SELECT Floor FROM sDyna")
            liste = curs.fetchall()
            if len(liste) != 0:
                for j in range(len(liste)):
                    for i in liste[j]:
                        if i == int(_cm_Floor):
                            checkUnique = 0

            if checkUnique == 1:
                curs.execute("INSERT INTO sDyna (Floor, Mass, Rigidity) VALUES (?,?,?)", (_cm_Floor,_lne_Mass,_lne_Rigidity))
                conn.commit()
                makeList()
                ui.lne_Mass.setEnabled(False)
                ui.lne_Rigidity.setEnabled(False)

            elif checkUnique == 0:
                QMessageBox.about(WinMain,"Error","This floor has added already. Use Change button to change")
                WinMain.show()

            checkUnique = 0

        else:
            ui.statusbar.showMessage("Error: Data must be entered.",10000)


    except Exception:
        QMessageBox.about(WinMain,'Error','Input can only be a number')
        pass

    
#---------------LIST---------------#
#----------------------------------#
def makeList():
    ui.tb_data.clearContents() #tablo içeriğini siler.
    curs.execute("SELECT * FROM sDyna")
    for rowIndex, rowData in enumerate(curs):
        for columnIndex, columnData in enumerate(rowData):
            ui.tb_data.setItem(rowIndex,columnIndex,QTableWidgetItem(str(columnData)))
    ui.lne_Mass.clear()
    ui.lne_Rigidity.clear()
    ui.cm_Floor.setCurrentIndex(-1) # "-1" comboBox içine hiçbirşey yazmamayı ifade eder.

    #-----Current Floor Number-----#
    curs.execute("SELECT COUNT (*) FROM sDyna")
    floorNumber = curs.fetchone()
    ui.label_savedFloor.setText(str(floorNumber[0]))

    # floorNumber[0] bizim kat sayımızı veriyor!!
    # floorNumbers = curs.fetchall()
    # ui.label_savedFloor.setText(str(len(floorNumbers)))



#------------RESET ALL---------------#
#------------------------------------#
def deleteAll():
    answer1 = QMessageBox.question(WinMain,"Delete all","Are you sure to reset data?",\
                                    QMessageBox.Yes | QMessageBox.No)
    if answer1 == QMessageBox.Yes:    
        curs.execute("DELETE FROM sDyna")
        conn.commit()
        ui.tb_data.clearContents() #tablo içeriğini siler.
        ui.lne_Mass.clear()
        ui.lne_Rigidity.clear()
        ui.cm_Floor.setCurrentIndex(-1) # "-1" comboBox içine hiçbirşey yazmamayı ifade eder.
        ui.lne_EQData.clear()
        ui.lne_Interval.clear()
        ui.label_savedFloor.clear()
    else:
        WinMain.show()


#------------------EXIT-------------------#
#-----------------------------------------#
def exit_():
    answer = QMessageBox.question(WinMain,"EXIT","Are you sure to exit from program?",\
                                    QMessageBox.Yes | QMessageBox.No)
    if answer == QMessageBox.Yes:
        conn.close()
        sys.exit(Application.exec_())
    else:
        WinMain.show()

#------------------DELETE-ROW-------------------#
#-----------------------------------------------#
def deleteRow():
    answer = QMessageBox.question(WinMain,"Delete the floor","Are you sure to delete this floor?",\
                                    QMessageBox.Yes | QMessageBox.No)

    if answer == QMessageBox.Yes:
        slcted = ui.tb_data.selectedItems()

        if bool(slcted)==True:
            dlt = slcted[0].text()

            try:
                curs.execute("DELETE FROM sDyna WHERE Floor='%s'"%(dlt))
                conn.commit()
                makeList()
                ui.statusbar.showMessage(str(dlt)+". floor's data have been deleted successfully.",10000) #10000 milisaniye=10 saniye mesaj görünecek.
            
            except Exception as Error:
                ui.statusbar.showMessage("Error:"+str(Error),10000)
        
        else:
            ui.statusbar.showMessage("There is no selected items.",10000)


    else:
        ui.statusbar.showMessage("Deleting process has been cancelled.",10000)
        WinMain.show()

#------------------CHANGE-ROW-------------------#
#-----------------------------------------------#
def changeRow():
    answer2 = QMessageBox.question(WinMain,"Change the floor's features","Are you sure to change this floor?",\
                                    QMessageBox.Yes | QMessageBox.No)
    uniquenumber=1            
    _lne_Mass = ui.lne_Mass.text()
    _lne_Rigidity = ui.lne_Rigidity.text()
    _cm_Floor = ui.cm_Floor.currentText()

    try:
        _lne_Mass=float(_lne_Mass)
        _lne_Rigidity=float(_lne_Rigidity)

        if answer2 == QMessageBox.Yes:

            if bool(_lne_Mass) and bool(_lne_Rigidity) and bool(_cm_Floor)==True:
                
                curs.execute("SELECT Floor FROM sDyna")
                liste = curs.fetchall()
                for j in range(len(liste)):
                    for i in liste[j]:
                        if i == int(_cm_Floor):
                            uniquenumber=0
                
                if uniquenumber==0:
                    curs.execute("UPDATE sDyna SET Mass = ?, Rigidity = ? WHERE Floor = ?" , (_lne_Mass,_lne_Rigidity,_cm_Floor))
                    conn.commit()
                    makeList()
                    ui.lne_Mass.setEnabled(False)
                    ui.lne_Rigidity.setEnabled(False)

                else:
                    QMessageBox.about(WinMain,"Error","Choosen floor cannot be find in table.Please save the floor first.")
                    # WinMain.show()

                
                checkUnique = 0

            else:
                ui.statusbar.showMessage("Error: All data must be entered.",10000)
        
        else:
            ui.statusbar.showMessage("Changing process has been cancelled.",10000)
            WinMain.show()
    
    except Exception:
        QMessageBox.about(WinMain,'Error','Input can only be a number')
        pass


#--------------------SEARCH---------------------#
#-----------------------------------------------#
def search_():
    search1=ui.cm_Floor.currentText()
    search2=ui.lne_Mass.text()
    search3=ui.lne_Rigidity.text()
    curs.execute("SELECT * FROM sDyna WHERE Floor=? OR Mass=? OR Rigidity=? OR\
                    (Mass=? AND Rigidity=?) OR (Floor=? AND Rigidity=?) OR (Mass=? AND Floor=?)",\
                    (search1,search2,search3,search2,search3,search1,search3,search2,search1))
    conn.commit
    ui.tb_data.clearContents()
    for rowIndex, rowData in enumerate(curs):
        for columnIndex, columnData in enumerate(rowData):
            ui.tb_data.setItem(rowIndex,columnIndex,QTableWidgetItem(str(columnData)))


#-----------------ABOUT-------------------#
#-----------------------------------------#
def about_():
    WinAbout.show()

#-----------------Manual-------------------#
#-----------------------------------------#
def manual_():
    WinManual.show()

#---------------EQ FILE-------------------#
#-----------------------------------------#
def eqfile():
    options = QFileDialog.Options()
    options |= QFileDialog.DontUseNativeDialog
    fileName, _ = QFileDialog.getOpenFileName(WinMain,"Select Earthquake File", "","All Files(*);; Text Files (*.txt);; Excel Files (*.xlsx);; CSV Files (*.csv)", options=options)
    ui.lne_EQData.setText(fileName)

#---------------SAVE EXCEL FILE-----------#
#-----------------------------------------#
def saveExcelFile():
    options = QFileDialog.Options()
    options |= QFileDialog.DontUseNativeDialog
    savedfileName, _ = QFileDialog.getSaveFileName(WinMain,"Save File","","Xlsx Files (*.xlsx)", options=options)
    if savedfileName !='':
        if savedfileName.endswith(".xlsx")==True:
            workbook=xlsxwriter.Workbook(savedfileName)
        else:

            workbook=xlsxwriter.Workbook(savedfileName+ ".xlsx")
        worksheet=workbook.add_worksheet("sDyna")
        curs.execute("SELECT * FROM sDyna")
        colnames = [desc[0] for desc in curs.description]
        data = curs.fetchall()


        # Start from the first cell. Rows and columns are zero indexed.
        row = 1
        col = 0
        row1=0

        # Iterate over the data and write it out row by row.
        
        for header in range(0,len(colnames)):
            worksheet.write(row1,col+header,colnames[header])
        

        for Floor, Mass, Rigidity in (data):
            worksheet.write(row, col,     Floor)
            worksheet.write(row, col + 1, Mass)
            worksheet.write(row, col + 2, Rigidity)
            row += 1

        workbook.close()


#-------------OPEN EXCEL FILE-------------#
#-----------------------------------------#
def openExcelFile():

    # Choosing Excel File
    options = QFileDialog.Options()
    options |= QFileDialog.DontUseNativeDialog
    fileName, _ = QFileDialog.getOpenFileName(WinMain,"Select Database", "","Excel Files (*.xlsx)", options=options)
    
    try: 
        
        wb = xlrd.open_workbook(fileName) 
        sheet = wb.sheet_by_index(0) 
        sheet.cell_value(0, 0)
        deleteAll()
        
        
        for i in range(1,sheet.nrows): 
            floorExcel=sheet.cell_value(i, 0)
            massExcel=sheet.cell_value(i,1)
            rigidityExcel=sheet.cell_value(i,2)
            curs.execute("INSERT INTO sDyna (Floor,Mass,Rigidity) VALUES (?,?,?)", (floorExcel,massExcel,rigidityExcel))
            conn.commit()
        
        makeList()
    
    except Exception:
        pass


#-------------------RUN-------------------#
#-----------------------------------------#
def run():
    
    curs.execute("Select Mass From sDyna")
    massList = curs.fetchall()
    for i in range(len(massList)):
        massList[i] = massList[i][0]

    curs.execute("Select Rigidity From sDyna")
    rigidityList = curs.fetchall()
    for i in range(len(rigidityList)):
        rigidityList[i] = rigidityList[i][0]
    
    #-----Total Floor Number-----#
    curs.execute("SELECT COUNT (*) FROM sDyna")
    floorNumber = curs.fetchone()

    _lne_EQData = ui.lne_EQData.text()
    _lne_Interval = ui.lne_Interval.text()

    #-----------Run MDOF-----------#
    yapi=Yapi(massList, rigidityList, floorNumber[0])

    completed = 0
    while completed < 12:
        completed += 0.0001
        ui.progressBar.setValue(completed)

    ui.txt_Results.setText("RESULTS")
    yapi.massMatrix()
    ui.txt_Results.append("\n" + "Mass Matrix\n" + (str(yapi.m_matrix)))
    yapi.rigidityMatrix()
    ui.txt_Results.append("\n" + "Stiffness Matrix\n" + (str(yapi.k_matrix)))
    yapi.naturalFrequency()
    yapi.dampingRatio(0.05)
    yapi.dampingMatrix()
    ui.txt_Results.append("\n" + "Damping Matrix\n" + (str(yapi.c_matrix)))
    ui.txt_Results.append("\n" + "Natural Frequency\n" + (str(yapi.wn)))
    ui.txt_Results.append("\n" + "Natural Period\n" + (str(yapi.Tn)))
    yapi.amplitudeCalc()
    ui.txt_Results.append("\n" + "Amplitudes\n" + (str(yapi.amp)))
    while completed <44:
        completed += 0.0001
        ui.progressBar.setValue(completed)
    yapi.generalMassMat()
    ui.txt_Results.append("\n" + "Generalized Mass Matrix\n" + (str(yapi.M_Generalized)))
    yapi.generalStiffnessMat()
    ui.txt_Results.append("\n" + "Generalized Stiffness Matrix\n" + (str(yapi.K_Generalized)))
    yapi.generalDampingMat()
    ui.txt_Results.append("\n" + "Generalized Damping Matrix\n" + (str(yapi.C_Generalized)))
    yapi.modeParticipatingFactor()
    ui.txt_Results.append("\n" + "Mode Participating Factor\n" + (str(yapi.lam)))

    while completed < 72:
        completed += 0.0001
        ui.progressBar.setValue(completed)
    
    yapi.effectiveParticipatingMass()
    ui.txt_Results.append("\n" + "Effective Participating Mass\n" + (str(yapi.M_eff)))

    while completed < 98:
        completed += 0.0001
        ui.progressBar.setValue(completed)

    yapi.earthquakeData(_lne_EQData,float(_lne_Interval))
    os.remove("EarthquakeData.png")
    yapi.spectra1()
    yapi.psuedoAcceleration()
    yapi.baseShear()
    ui.txt_Results.append("\n" + "Modes' Base Shear forces\n")
    for i in range(0,yapi.storeynumber):
        ui.txt_Results.append(str(round(yapi.Ft[i][0],2)))
    yapi.baseShearSRSS()
    ui.txt_Results.append("\n" + "Total Base Shear Force with SRSS = " + str(round(yapi.totalFt[0],2)) + " kN")

    while completed < 100:
        completed += 0.0001
        ui.progressBar.setValue(completed)
    
    ui.progressBar.hide()
    ui.progressBar.setProperty("value", 0)
    



#--------------ProgressBar----------------#
#-----------------------------------------#
def run_():
    ui.progressBar.setFormat("Running %p.00%")
    ui.progressBar.show()

    answer4 = QMessageBox.question(WinMain,"Run","Are you sure to run?",\
                                    QMessageBox.Yes | QMessageBox.No)
    if answer4 == QMessageBox.Yes:
        try:
            run()
        except Exception as Error:
            ui.txt_Results.clear()
            ui.statusbar.showMessage("Error: There is an invalid value",10000)
    else:
        ui.progressBar.hide()
        ui.statusbar.showMessage("Running has been cancelled.",10000)
        WinMain.show()
    

#-------------WORD--------------#
#-------------------------------#
def word():

    from docx import Document
    from docx.shared import Inches, Pt
    import numpy as np
    import matplotlib.pyplot as plt
    import matplotlib
    from docx.enum.table import WD_TABLE_ALIGNMENT

    options = QFileDialog.Options()
    options |= QFileDialog.DontUseNativeDialog
    savedfileName, _ = QFileDialog.getSaveFileName(WinMain,"Save File","","Docx Files (*.docx)", options=options)
    if savedfileName !='':
        if savedfileName.endswith(".docx")==True:
            savedfileName=savedfileName
        else:

            savedfileName=(savedfileName+ ".docx")

        completed = 0
        while completed < 12:
            completed += 0.0001
            ui.progressBar.setValue(completed)

        #-------Run bof getting results------#
        curs.execute("Select Mass From sDyna")
        massList = curs.fetchall()
        for i in range(len(massList)):
            massList[i] = massList[i][0]
        curs.execute("Select Rigidity From sDyna")
        rigidityList = curs.fetchall()
        for i in range(len(rigidityList)):
            rigidityList[i] = rigidityList[i][0]
        curs.execute("SELECT COUNT (*) FROM sDyna")
        floorNumber = curs.fetchone()
        _lne_EQData = ui.lne_EQData.text()
        _lne_Interval = ui.lne_Interval.text()
        yapi=Yapi(massList, rigidityList, floorNumber[0])
        yapi.rigidityMatrix()
        yapi.massMatrix()
        yapi.naturalFrequency()
        yapi.dampingRatio(0.05)
        yapi.dampingMatrix()
        yapi.amplitudeCalc()
        yapi.ModalShapes()
        yapi.generalMassMat()
        yapi.generalStiffnessMat()
        yapi.generalDampingMat()
        yapi.modeParticipatingFactor()  
        yapi.effectiveParticipatingMass()
        yapi.earthquakeData(_lne_EQData,float(_lne_Interval))
        yapi.spectra2()
        yapi.spectra1()
        yapi.psuedoAcceleration()
        yapi.baseShear()
        yapi.baseShearSRSS()
        


        #-----Create and edit .docx file-----#
        document=Document()
        header=document.add_heading("sDyna REPORT",0)
        header.alignment=1
        document.add_paragraph("This report has been automatically generated by sDyna software. ")
        
        #--------Writing down Matrixes-------#
        curs.execute("SELECT COUNT (*) FROM sDyna")
        rowNumber = curs.fetchone()
        ArrayNameLocation=math.floor(rowNumber[0]/2)

        #--------Mass Matrix--------#
        document.add_heading("Mass Matrix\n",level=1)
        tableMass = document.add_table(rows=rowNumber[0], cols=rowNumber[0]+1)
        tableMass.allow_autofit = False
        # table.style = "Table Grid"
        for i in range(1,rowNumber[0]+1):
            for cell in tableMass.columns[i].cells:
                cell.width = Inches(0.5)

        #----------Labeling Mass Matrix-----------#
        cell = tableMass.cell(ArrayNameLocation, 0)
        cell.width = Inches(1.5)
        cell.text = 'm ='
        
        #----------Creating Mass Matrix-----------#
        for i in range(0,rowNumber[0]):
            for j in range(0,rowNumber[0]):
                cell=tableMass.cell(i,j+1)
                cell.text="{}".format(yapi.m_matrix[i][j])
        
        for row in tableMass.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)

        #----------Labeling Rigidity Matrix-----------#
        document.add_heading("Stiffness Matrix\n",level=1)
        tableRigidity = document.add_table(rows=rowNumber[0], cols=rowNumber[0]+1)
        tableRigidity.allow_autofit = False
        for i in range(1,rowNumber[0]+1):
            for cell in tableRigidity.columns[i].cells:
                cell.width = Inches(1.1)
        cell = tableRigidity.cell(ArrayNameLocation, 0)
        cell.width = Inches(1.8)
        cell.text = 'k ='
        
        #----------Creating Rigidity Matrix-----------#
        
        for i in range(0,rowNumber[0]):
            for j in range(0,rowNumber[0]):
                cell=tableRigidity.cell(i,j+1)
                cell.text="{}".format(round(yapi.k_matrix[i][j],2))
        
        for row in tableRigidity.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)

        #----------Labeling Damping Matrix-----------#
        document.add_heading("Damping Matrix\n",level=1)
        document.add_paragraph("Damping ratio has been taken as 0.05.")
        tableDamping = document.add_table(rows=rowNumber[0], cols=rowNumber[0]+1)
        tableDamping.allow_autofit = False
        for i in range(1,rowNumber[0]+1):
            for cell in tableDamping.columns[i].cells:
                cell.width = Inches(1.1)
        cell = tableDamping.cell(ArrayNameLocation, 0)
        cell.width = Inches(2)
        cell.text = 'c ='
        
        #----------Creating Damping Matrix-----------#
        
        for i in range(0,rowNumber[0]):
            for j in range(0,rowNumber[0]):
                cell=tableDamping.cell(i,j+1)
                cell.text="{}".format(round(yapi.c_matrix[i][j],2))
        
        for row in tableDamping.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)
        
        #----------Creating Natural Frequency-----------#
        document.add_heading("Natural Frequencies and Periods",level=1)
        document.add_paragraph("")
        for i in range(0,rowNumber[0]):
            p=document.add_paragraph("")
            a = p.add_run("ω")
            font = a.font
            font.size = Pt(9)
            sub_p=p.add_run("{}".format(i+1))
            sub_p.font.subscript=True
            font = sub_p.font
            font.size = Pt(9)
            b = p.add_run(" = {} rad/sec           --------------->     T".format(round(yapi.wn[i][0],3)))
            font = b.font
            font.size = Pt(9)
            sub_p=p.add_run("{}".format(i+1))
            font = sub_p.font
            font.size = Pt(9)
            c = p.add_run("= {} sec".format(round(yapi.Tn[i][0],3)))
            font = c.font
            font.size = Pt(9)
            sub_p.font.subscript=True
        
        
        #----------Amplitudes-----------#
        document.add_page_break()
        document.add_heading("Modes' Amplitudes",level=1 )
        document.add_paragraph("")

        tableAmplitude=document.add_table(rows=rowNumber[0],cols=rowNumber[0]+1)

        for i in range(0,rowNumber[0]):
            cell1=tableAmplitude.cell(i,0)
            a1=cell1.add_paragraph("Mode {} Amplitudes".format(i+1))
            for j in range(0,rowNumber[0]):
                cell=tableAmplitude.cell(i,j+1)
                a=cell.add_paragraph("φ")
                sub_text=a.add_run("{}{}".format(i+1,j+1))
                sub_text.font.subscript=True
                a.add_run(" = {}".format(round(yapi.amp[i][j],3)))
                
                tableAmplitude.allow_autofit=False
                tableAmplitude.style = "Table Grid"

        for row in tableAmplitude.rows:
            row.height = Inches(0.4)
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)

        #----------Mode Shapes-----------#
        document.add_paragraph("")
        document.add_picture("ModeShapes.png",width=Inches(6),height=Inches(3))
        os.remove("ModeShapes.png")

        document.add_page_break()

        #----------Generalized Mass Matrix-----------#
        document.add_heading("Generalized Mass Matrix",level=1)
        document.add_paragraph("")
        tableGenMass = document.add_table(rows=rowNumber[0], cols=rowNumber[0]+1)
        tableGenMass.allow_autofit = False
        for i in range(1,rowNumber[0]+1):
            for cell in tableGenMass.columns[i].cells:
                cell.width = Inches(0.5)
        cell = tableGenMass.cell(ArrayNameLocation, 0)
        cell.width = Inches(1.5)
        cell.text = 'M ='
        
        #----------Creating Generalized Mass Matrix-----------#
        
        for i in range(0,rowNumber[0]):
            for j in range(0,rowNumber[0]):
                cell=tableGenMass.cell(i,j+1)
                cell.text="{}".format(round(yapi.M_Generalized[i][j],2))
        
        for row in tableGenMass.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)

    #----------Generalized Stiffness Matrix-----------#
        document.add_heading("Generalized Stiffness Matrix",level=1)
        document.add_paragraph("")
        tableGenStiffness = document.add_table(rows=rowNumber[0], cols=rowNumber[0]+1)
        tableGenStiffness.allow_autofit = False
        for i in range(1,rowNumber[0]+1):
            for cell in tableGenStiffness.columns[i].cells:
                cell.width = Inches(0.5)
        cell = tableGenStiffness.cell(ArrayNameLocation, 0)
        cell.width = Inches(1.5)
        cell.text = 'K ='
        
        #----------Creating Generalized Stiffness Matrix-----------#
        
        for i in range(0,rowNumber[0]):
            for j in range(0,rowNumber[0]):
                cell=tableGenStiffness.cell(i,j+1)
                cell.text="{}".format(round(yapi.K_Generalized[i][j],2))
        
        for row in tableGenStiffness.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)

        #----------Generalized Damping Matrix-----------#
        document.add_heading("Generalized Damping Matrix",level=1)
        document.add_paragraph("")
        tableGenDamping = document.add_table(rows=rowNumber[0], cols=rowNumber[0]+1)
        tableGenDamping.allow_autofit = False
        for i in range(1,rowNumber[0]+1):
            for cell in tableGenDamping.columns[i].cells:
                cell.width = Inches(0.5)
        cell = tableGenDamping.cell(ArrayNameLocation, 0)
        cell.width = Inches(1.5)
        cell.text = 'C ='
        
        #----------Creating Generalized Damping Matrix-----------#
        
        for i in range(0,rowNumber[0]):
            for j in range(0,rowNumber[0]):
                cell=tableGenDamping.cell(i,j+1)
                cell.text="{}".format(round(yapi.C_Generalized[i][j],2))
        
        for row in tableGenDamping.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)

        while completed < 58:
            completed += 0.0001
            ui.progressBar.setValue(completed)

        #----------MODAL PARTICIPATING FACTOR FOR EACH MODES-----------#
        document.add_heading("Modal Participating Factors",level=1)
        document.add_paragraph("")
        for i in range(0,rowNumber[0]):
            p=document.add_paragraph("")
            a = p.add_run("Γ")
            font = a.font
            font.size = Pt(9)
            sub_p=p.add_run("x{}".format(i+1))
            sub_p.font.subscript=True
            font = sub_p.font
            font.size = Pt(9)
            b = p.add_run(" = {} ".format(round(yapi.lam[i][0],3)))
            font = b.font
            font.size = Pt(9)
            sub_p.font.subscript=True
        
        #----------Effective Participating Mass Of Each Mode-----------#
        document.add_heading("Effective Participating Mass of Each Mode",level=1)
        document.add_paragraph("")
        for i in range(0,rowNumber[0]):
            p=document.add_paragraph("")
            a = p.add_run("M")
            font = a.font
            font.size = Pt(9)
            sub_p=p.add_run("x{}".format(i+1))
            sub_p.font.subscript=True
            font = sub_p.font
            font.size = Pt(9)
            b = p.add_run(" = {} ".format(round(yapi.M_eff[i][0],3)))
            font = b.font
            font.size = Pt(9)
            sub_p.font.subscript=True
        
        #----------Earthquake Acceleration Data-----------#
        
        document.add_picture("EarthquakeData.png",width=Inches(6),height=Inches(4))
        os.remove("EarthquakeData.png")
        document.add_paragraph("")

        #----------Psuedo Displacement Respond Spectrum-----------#
        
        document.add_picture("PseudoDisplacement.png",width=Inches(6),height=Inches(4))
        os.remove("PseudoDisplacement.png")
        document.add_paragraph("")
        document.add_picture("PseudoAcceleration.png",width=Inches(6),height=Inches(4))
        os.remove("PseudoAcceleration.png")


        #----------Psuedo Acceleration Respond Spectrum-----------#
        document.add_heading("Psuedo Acceleration Respond Spectrum",level=1)
        document.add_paragraph("")
        for i in range(0,rowNumber[0]):
            p=document.add_paragraph("")
            a = p.add_run("S")
            font = a.font
            font.size = Pt(11)
            sub_p=p.add_run("ae{}".format(i+1))
            sub_p.font.subscript=True
            font = sub_p.font
            font.size = Pt(11)
            b = p.add_run(" = {} ".format(round(yapi.Sae[i][0],3)))
            font = b.font
            font.size = Pt(11)
            sub_p.font.subscript=True
        
        #----------Mode's Base Shear Forces-----------#
        document.add_heading("Modes' Base Shear Forces",level=1)
        document.add_paragraph("")
        for i in range(0,rowNumber[0]):
            p=document.add_paragraph("")
            a = p.add_run("V")
            font = a.font
            font.size = Pt(11)
            sub_p=p.add_run("x{}".format(i+1))
            sub_p.font.subscript=True
            font = sub_p.font
            font.size = Pt(11)
            b = p.add_run(" = {} ".format(round(yapi.Ft[i][0],3)))
            font = b.font
            font.size = Pt(11)
            sub_p.font.subscript=True
        
        #----------Estimation of Base Shear Force with SRSS method-----------#
        document.add_heading("Total Base Shear Force with SRSS method",level=1)
        document.add_paragraph("")
        p=document.add_paragraph("")
        a = p.add_run("V")
        font = a.font
        font.size = Pt(11)
        sub_p=p.add_run("T")
        sub_p.font.subscript=True
        font = sub_p.font
        font.size = Pt(11)
        b = p.add_run(" = {} ".format(round(yapi.totalFt[0],3)))
        font = b.font
        font.size = Pt(11)
        sub_p.font.subscript=True


        while completed < 100:
            completed += 0.0001
            ui.progressBar.setValue(completed)
        
        ui.progressBar.hide()
        ui.progressBar.setProperty("value", 0)

        document.save(savedfileName)
    else:
        ui.progressBar.hide()

#--------------Printing Bar----------------#
#-----------------------------------------#
def word_():
    ui.progressBar.setFormat("Printing %p.00%")
    ui.progressBar.show()

    answer4 = QMessageBox.question(WinMain,"Print","Are you sure to print?",\
                                    QMessageBox.Yes | QMessageBox.No)
    if answer4 == QMessageBox.Yes:
        try:
            word()
        except Exception as Error:
            os.remove(savedfileName)
            ui.statusbar.showMessage("Error: There is an invalid value",10000)

    else:
        ui.progressBar.hide()
        ui.statusbar.showMessage("Printing process has been cancelled.",10000)
        WinMain.show()
    
#-----------------Links-------------------#
#-----------------------------------------#
def linkOnur():
    import webbrowser
    webbrowser.open('https://www.linkedin.com/in/onur-%C3%BClk%C3%BC-836434128/')

def linkAli():
    import webbrowser
    webbrowser.open('https://www.linkedin.com/in/alitatici/')

def linkAFAD():
    import webbrowser
    webbrowser.open('https://tadas.afad.gov.tr/')


#---------------SIGNAL-SLOT---------------#
#-----------------------------------------#
ui.pb_Add.clicked.connect(addData)
ui.pb_reset.clicked.connect(deleteAll)
ui.pb_Exit.clicked.connect(exit_)
ui.pb_dltRow.clicked.connect(deleteRow)
ui.pb_find.clicked.connect(search_)
ui.pb_list.clicked.connect(makeList)
ui.actionAbout.triggered.connect(about_)
ui.actionManual_2.triggered.connect(manual_)
ui.pb_change.clicked.connect(changeRow)
ui.cm_Floor.activated.connect(comboact)
ui.pb_fromfile.clicked.connect(eqfile)
ui.pb_Save.clicked.connect(saveExcelFile)
ui.pb_Open.clicked.connect(openExcelFile)
ui.pb_run.clicked.connect(run_)
ui.pb_Print.clicked.connect(word_)
uiAbout.pushButton_2.clicked.connect(linkOnur)
uiAbout.pushButton.clicked.connect(linkAli)
uiManual.commandLinkButton.clicked.connect(linkAFAD)

sys.exit(Application.exec_())

